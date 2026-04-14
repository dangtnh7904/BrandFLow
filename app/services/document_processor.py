"""
DocumentIngestor — Multi-Format Smart Parser
=============================================
Hỗ trợ định dạng:
  • PDF   — pdfplumber (text layer) → pdfminer fallback → LlamaParse Cloud (nếu có API key)
  • DOCX  — python-docx (paragraphs + tables + headers + footers)
  • TXT / MD — auto-detect encoding (utf-8 → latin-1 → cp1252)
  • CSV   — pandas (convert to human-readable text)
  • XLSX / XLS — pandas (sheet-by-sheet, table format)
  • HTML  — BeautifulSoup text extraction
"""

import io
import os
import re
from dataclasses import dataclass, field
from datetime import datetime
from typing import List, Optional

from dotenv import load_dotenv

load_dotenv()

from langchain_core.documents import Document
from langchain_chroma import Chroma
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter


# ─── Kết quả parse từng file ─────────────────────────────────────────────────

@dataclass
class FileParseResult:
    filename: str
    ext: str
    success: bool
    text: str = ""
    char_count: int = 0
    page_count: int = 0
    sheet_count: int = 0
    method: str = ""
    error: Optional[str] = None

    def __post_init__(self):
        self.char_count = len(self.text)


# ─── Helpers ─────────────────────────────────────────────────────────────────

def _try_read_text_file(path: str) -> str:
    """Đọc file text với auto-detect encoding (utf-8 → latin-1 → cp1252)."""
    for enc in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            with open(path, "r", encoding=enc, errors="strict") as f:
                return f.read()
        except (UnicodeDecodeError, LookupError):
            continue
    # Last resort: ignore errors
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def _parse_pdf(path: str) -> FileParseResult:
    """
    PDF parser theo 3 tầng ưu tiên:
      1. pdfplumber  — nhanh, chính xác với PDF có text layer
      2. pdfminer    — fallback mạnh hơn cho layout phức tạp
      3. LlamaParse  — AI cloud parser cho PDF scan (cần API key)
    """
    text = ""
    page_count = 0
    method = ""

    # Tầng 1: pdfplumber
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            page_count = len(pdf.pages)
            lines = []
            for page in pdf.pages:
                # Extract text
                page_text = page.extract_text(x_tolerance=2, y_tolerance=2)
                if page_text:
                    lines.append(page_text)
                # Extract tables
                tables = page.extract_tables()
                for table in tables:
                    if not table:
                        continue
                    for row in table:
                        row_text = " | ".join(
                            cell.strip() if cell else "" for cell in row
                        )
                        if row_text.strip(" |"):
                            lines.append(row_text)
            text = "\n".join(lines)
            method = "pdfplumber"
        print(f"   ✅ [PDF/pdfplumber] {page_count} trang, {len(text)} ký tự")
    except Exception as e:
        print(f"   ⚠️ pdfplumber lỗi: {e}")

    # Tầng 2: pdfminer (nếu pdfplumber ra rỗng)
    if not text.strip():
        try:
            from pdfminer.high_level import extract_text as pdfminer_extract
            text = pdfminer_extract(path)
            method = "pdfminer"
            print(f"   ✅ [PDF/pdfminer] {len(text)} ký tự")
        except Exception as e:
            print(f"   ⚠️ pdfminer lỗi: {e}")

    # Tầng 3: LlamaParse AI Cloud (nếu vẫn rỗng và có API key)
    if not text.strip():
        api_key = os.getenv("LLAMA_CLOUD_API_KEY", "").strip()
        if api_key:
            try:
                from llama_parse import LlamaParse
                print(f"   🚀 Gọi LlamaParse Cloud AI...")
                parser = LlamaParse(api_key=api_key, result_type="markdown", verbose=False)
                parsed_docs = parser.load_data(path)
                text = "\n".join(d.text for d in parsed_docs)
                method = "llamaparse_cloud"
                print(f"   ✅ [LlamaParse] {len(text)} ký tự")
            except Exception as e:
                print(f"   ⚠️ LlamaParse lỗi: {e}")

    if not text.strip():
        return FileParseResult(
            filename=os.path.basename(path), ext="pdf", success=False,
            error="Không thể trích xuất nội dung từ PDF này (có thể là ảnh scan không có OCR).",
            page_count=page_count, method=method or "none",
        )

    return FileParseResult(
        filename=os.path.basename(path), ext="pdf", success=True,
        text=text, page_count=page_count, method=method,
    )


def _parse_docx(path: str) -> FileParseResult:
    """
    DOCX parser đầy đủ:
      - Paragraphs (text thường + heading)
      - Tables (format: col1 | col2 | col3)
      - Headers & Footers
    """
    try:
        import docx as python_docx
        doc = python_docx.Document(path)
        sections: List[str] = []

        # Header
        for section in doc.sections:
            if section.header and section.header.paragraphs:
                header_text = " | ".join(
                    p.text.strip() for p in section.header.paragraphs if p.text.strip()
                )
                if header_text:
                    sections.append(f"[HEADER] {header_text}")

        # Body: paragraphs + tables (preserving order via document XML)
        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

            if tag == "p":
                # Paragraph
                para = python_docx.text.paragraph.Paragraph(element, doc)
                text = para.text.strip()
                if text:
                    # Thêm prefix nếu là heading
                    style = para.style.name if para.style else ""
                    if "Heading" in style:
                        level = style.replace("Heading", "").strip()
                        prefix = "#" * int(level) if level.isdigit() else "##"
                        sections.append(f"{prefix} {text}")
                    else:
                        sections.append(text)

            elif tag == "tbl":
                # Table
                tbl = python_docx.table.Table(element, doc)
                try:
                    rows = []
                    for i, row in enumerate(tbl.rows):
                        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                        rows.append(" | ".join(cells))
                        if i == 0:
                            # Separator dưới header row
                            rows.append(" | ".join(["---"] * len(cells)))
                    sections.append("\n".join(rows))
                except Exception:
                    pass

        # Footer
        for section in doc.sections:
            if section.footer and section.footer.paragraphs:
                footer_text = " | ".join(
                    p.text.strip() for p in section.footer.paragraphs if p.text.strip()
                )
                if footer_text:
                    sections.append(f"[FOOTER] {footer_text}")

        text = "\n\n".join(sections)
        print(f"   ✅ [DOCX] {len(doc.paragraphs)} đoạn, {len(doc.tables)} bảng, {len(text)} ký tự")
        return FileParseResult(
            filename=os.path.basename(path), ext="docx", success=True,
            text=text, method="python-docx",
        )
    except Exception as e:
        return FileParseResult(
            filename=os.path.basename(path), ext="docx", success=False,
            error=str(e), method="python-docx",
        )


def _parse_txt_md(path: str, ext: str) -> FileParseResult:
    """TXT / MD / LOG — auto-detect encoding."""
    try:
        text = _try_read_text_file(path)
        print(f"   ✅ [{ext.upper()}] {len(text)} ký tự")
        return FileParseResult(
            filename=os.path.basename(path), ext=ext, success=True,
            text=text, method="plain-text",
        )
    except Exception as e:
        return FileParseResult(
            filename=os.path.basename(path), ext=ext, success=False,
            error=str(e), method="plain-text",
        )


def _parse_csv(path: str) -> FileParseResult:
    """CSV — đọc bằng pandas, convert sang dạng text bảng."""
    try:
        import pandas as pd
        # Thử nhiều encoding
        df = None
        for enc in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
            try:
                df = pd.read_csv(path, encoding=enc)
                break
            except Exception:
                continue
        if df is None:
            raise ValueError("Không thể đọc CSV với các encoding phổ biến.")

        lines = []
        lines.append(f"[CSV File: {os.path.basename(path)}]")
        lines.append(f"Số dòng: {len(df)}, Số cột: {len(df.columns)}")
        lines.append(f"Các cột: {', '.join(str(c) for c in df.columns)}")
        lines.append("")

        # Chuyển thành text dạng bảng (tối đa 500 dòng để tránh quá lớn)
        lines.append(df.head(500).to_string(index=False))

        text = "\n".join(lines)
        print(f"   ✅ [CSV] {len(df)} dòng × {len(df.columns)} cột, {len(text)} ký tự")
        return FileParseResult(
            filename=os.path.basename(path), ext="csv", success=True,
            text=text, method="pandas-csv",
        )
    except Exception as e:
        return FileParseResult(
            filename=os.path.basename(path), ext="csv", success=False,
            error=str(e), method="pandas-csv",
        )


def _parse_excel(path: str, ext: str) -> FileParseResult:
    """XLSX / XLS — đọc từng sheet bằng pandas."""
    try:
        import pandas as pd
        xl = pd.ExcelFile(path)
        sheet_names = xl.sheet_names
        sections = [f"[Excel File: {os.path.basename(path)}]", f"Số sheet: {len(sheet_names)}"]

        for sheet_name in sheet_names:
            df = xl.parse(sheet_name)
            sections.append(f"\n--- Sheet: {sheet_name} ({len(df)} dòng × {len(df.columns)} cột) ---")
            sections.append(f"Các cột: {', '.join(str(c) for c in df.columns)}")
            sections.append(df.head(500).to_string(index=False))

        text = "\n".join(sections)
        print(f"   ✅ [Excel] {len(sheet_names)} sheet(s), {len(text)} ký tự")
        return FileParseResult(
            filename=os.path.basename(path), ext=ext, success=True,
            text=text, sheet_count=len(sheet_names), method="pandas-excel",
        )
    except Exception as e:
        return FileParseResult(
            filename=os.path.basename(path), ext=ext, success=False,
            error=str(e), method="pandas-excel",
        )


def _parse_html(path: str) -> FileParseResult:
    """HTML — BeautifulSoup plain text extraction."""
    try:
        from bs4 import BeautifulSoup
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            soup = BeautifulSoup(f.read(), "html.parser")
        # Xóa script + style
        for tag in soup(["script", "style"]):
            tag.decompose()
        text = soup.get_text(separator="\n")
        text = re.sub(r"\n{3,}", "\n\n", text)
        print(f"   ✅ [HTML] {len(text)} ký tự")
        return FileParseResult(
            filename=os.path.basename(path), ext="html", success=True,
            text=text, method="beautifulsoup",
        )
    except Exception as e:
        return FileParseResult(
            filename=os.path.basename(path), ext="html", success=False,
            error=str(e), method="beautifulsoup",
        )


# ─── Main Ingestor Class ──────────────────────────────────────────────────────

class DocumentIngestor:
    """
    Bộ đọc tài liệu đa định dạng cho BrandFlow.

    Định dạng hỗ trợ:
      PDF, DOCX, DOC, TXT, MD, LOG, CSV, XLSX, XLS, HTML, HTM
    """

    SUPPORTED_EXTENSIONS = {
        "pdf", "docx", "doc",
        "txt", "md", "log",
        "csv",
        "xlsx", "xls",
        "html", "htm",
    }

    def __init__(self, persist_directory: str = "./chroma_db", tenant_id: str = "default"):
        self.persist_directory = persist_directory
        self.tenant_id = tenant_id
        self.embeddings = GoogleGenerativeAIEmbeddings(model="models/gemini-embedding-001")

    # ── Public API ────────────────────────────────────────────────────────────

    def ingest_file(self, file_path: str, force_ai: bool = False) -> str:
        """
        Đọc file và trả về raw text.
        Để tương thích ngược với code cũ (chỉ cần text string).
        """
        result = self.parse_file(file_path)
        if not result.success:
            raise ValueError(result.error or f"Không đọc được file: {file_path}")
        return result.text

    def parse_file(self, file_path: str) -> FileParseResult:
        """
        Parse file và trả về FileParseResult (có metadata đầy đủ).
        """
        ext = os.path.splitext(file_path)[-1].lstrip(".").lower()
        fname = os.path.basename(file_path)
        print(f"📄 [Ingestion] Đang quét: {fname} (định dạng: .{ext})")

        if ext not in self.SUPPORTED_EXTENSIONS:
            return FileParseResult(
                filename=fname, ext=ext, success=False,
                error=(
                    f"Định dạng .{ext} chưa được hỗ trợ. "
                    f"Các định dạng hỗ trợ: {', '.join(sorted(self.SUPPORTED_EXTENSIONS))}"
                ),
            )

        if ext == "pdf":
            return _parse_pdf(file_path)
        elif ext in ("docx", "doc"):
            return _parse_docx(file_path)
        elif ext in ("txt", "md", "log"):
            return _parse_txt_md(file_path, ext)
        elif ext == "csv":
            return _parse_csv(file_path)
        elif ext in ("xlsx", "xls"):
            return _parse_excel(file_path, ext)
        elif ext in ("html", "htm"):
            return _parse_html(file_path)
        else:
            # Không bao giờ đến đây do check trên, nhưng để tránh lint warning
            return FileParseResult(filename=fname, ext=ext, success=False, error="Unsupported")

    def ingest_url(self, url: str) -> str:
        """Đọc nội dung text từ URL."""
        try:
            from langchain_community.document_loaders import WebBaseLoader
            loader = WebBaseLoader(url)
            docs = loader.load()
            return "\n".join([doc.page_content for doc in docs])
        except Exception as e:
            raise ValueError(f"Ingestion URL failed: {e}")

    def clean_text(self, text: str) -> str:
        """Làm sạch văn bản cơ bản."""
        if not text:
            return ""
        text = text.replace("\x00", "")
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    def semantic_chunking(self, text: str) -> List[str]:
        """Cắt văn bản thành các fragments tối ưu cho RAG."""
        print("✂️  [Chunking] Đang phân mảnh tài liệu...")
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=150,
            separators=["\n\n", "\n", ".", " ", ""],
        )
        chunks = splitter.split_text(text)
        print(f"   ✅ Đã phân thành {len(chunks)} chunks.")
        return chunks

    def process_and_store_text(self, raw_text: str, filename: str, category: str):
        """Pipeline: Cleansing → Chunking → VectorDB."""
        current_date = datetime.now().isoformat()
        try:
            cleaned_text = self.clean_text(raw_text)
            chunks = self.semantic_chunking(cleaned_text)
            if not chunks:
                print(f"   ⚠️ Không có chunk nào để lưu cho [{filename}].")
                return

            documents = [
                Document(
                    page_content=chunk,
                    metadata={
                        "source": filename,
                        "category": category,
                        "date": current_date,
                        "index": i,
                    },
                )
                for i, chunk in enumerate(chunks)
            ]

            print(f"💾 [Storage] Đang lưu {len(documents)} bản ghi vào ChromaDB...")
            collection_name = f"brandflow_memory_{self.tenant_id}"
            vectorstore = Chroma(
                collection_name=collection_name,
                embedding_function=self.embeddings,
                persist_directory=self.persist_directory,
            )
            vectorstore.add_documents(documents)
            print(f"🎉 Đã lưu thành công [{filename}] vào bộ não.")
        except Exception as e:
            print(f"🔴 [Storage Lỗi]: {e}")
            raise
