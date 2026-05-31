import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = docx.Document()

# Styles
styles = doc.styles
normal_style = styles['Normal']
normal_style.font.name = 'Arial'
normal_style.font.size = Pt(11)

# Title
title = doc.add_paragraph('BÁO CÁO DỰ ÁN CHI TIẾT & KẾ HOẠCH KINH DOANH')
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.size = Pt(24)
title.runs[0].font.bold = True
title.runs[0].font.color.rgb = RGBColor(0, 51, 102)

subtitle = doc.add_paragraph('NỀN TẢNG BRANDFLOW - AI MULTI-AGENT MARKETING SYSTEM')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(16)
subtitle.runs[0].font.bold = True
subtitle.runs[0].font.color.rgb = RGBColor(128, 128, 128)

doc.add_paragraph('\n')
doc.add_paragraph('Đơn vị thực hiện: Ban Quản Trị Dự Án BrandFlow').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Phiên bản tài liệu: 2.0 (Premium B2B Edition)').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Mức độ bảo mật: Lưu hành nội bộ & Đối tác Chiến lược').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

# Function to add heading
def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    h.runs[0].font.name = 'Arial'
    if level == 1:
        h.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        h.runs[0].font.bold = True
        h.runs[0].font.size = Pt(16)
    elif level == 2:
        h.runs[0].font.color.rgb = RGBColor(0, 102, 204)
        h.runs[0].font.bold = True
        h.runs[0].font.size = Pt(14)
    return h

# Content
add_heading('1. TÓM TẮT ĐIỀU HÀNH (EXECUTIVE SUMMARY)', 1)
doc.add_paragraph("BrandFlow là nền tảng B2B SaaS đột phá, ứng dụng cấu trúc Trí tuệ Nhân tạo Đa đặc vụ (AI Multi-Agent System) nhằm tái định nghĩa toàn bộ quy trình vận hành và thực thi Marketing của doanh nghiệp. Trong bối cảnh chuyển đổi số, khi chi phí nhân sự chuyên môn ngày càng đắt đỏ và các giải pháp AI đại trà như ChatGPT bộc lộ nhiều điểm yếu về tính rủi ro tài chính (Financial Hallucination), BrandFlow ra đời như một giải pháp toàn diện.")
doc.add_paragraph("Hệ sinh thái BrandFlow vận hành như một 'Phòng Marketing Ảo' với 5 Đặc vụ AI chuyên biệt, có khả năng tự động hóa 100% khối lượng công việc từ nghiên cứu thị trường, hoạch định chiến lược, phân bổ ngân sách, đến sáng tạo hình ảnh và nội dung. Đặc biệt, sự hiện diện của Math Engine - thuật toán quản trị rủi ro tài chính độc quyền, giúp đảm bảo tính an toàn tuyệt đối cho mọi khoản đầu tư quảng cáo của khách hàng.")
doc.add_paragraph("Với mô hình kinh doanh SaaS linh hoạt, dự án hướng tới mục tiêu chiếm lĩnh thị trường B2B Marketing Tech tại khu vực, duy trì tỷ lệ LTV:CAC lý tưởng ở mức 5.3:1 và chuyển đổi thành công 10% khách hàng từ bản miễn phí sang trả phí định kỳ.")

add_heading('2. PHÂN TÍCH VẤN ĐỀ VÀ THỰC TRẠNG THỊ TRƯỜNG (MARKET ANALYSIS)', 1)
add_heading('2.1. Nỗi đau của doanh nghiệp (Pain Points)', 2)
doc.add_paragraph("Thị trường hiện tại đang chứng kiến 3 điểm nghẽn nghiêm trọng cản trở quá trình tăng trưởng của các SME và Agency:", style='List Bullet')
doc.add_paragraph("Chi Phí Vận Hành Không Linh Hoạt: Quỹ lương cho một bộ phận Marketing tiêu chuẩn (Bao gồm CMO, Data Analyst, Designer, Copywriter) dao động từ 3,000 USD - 5,000 USD/tháng tại các quốc gia đang phát triển. Đây là gánh nặng lớn đối với nhóm doanh nghiệp SME.", style='List Bullet 2')
doc.add_paragraph("Quy Trình Thiếu Tính Hệ Thống & Chậm Trễ: Việc hoạch định chiến lược thủ công tốn trung bình 2-4 tuần, khiến doanh nghiệp chậm chân trong các chiến dịch bắt trend hoặc phản ứng với đối thủ cạnh tranh.", style='List Bullet 2')
doc.add_paragraph("Sự Khủng Hoảng Niềm Tin Vào AI Đại Trà: Việc lạm dụng các công cụ LLMs đa dụng (ChatGPT, Gemini) trong quản trị ngân sách thường xuyên dẫn đến các sai lầm nghiêm trọng do AI thiếu khả năng phân tích dữ liệu thị trường theo thời gian thực.", style='List Bullet 2')

add_heading('2.2. Quy mô và Tiềm năng thị trường (TAM, SAM, SOM)', 2)
doc.add_paragraph("Tổng quy mô thị trường (TAM): Thị trường Công nghệ Tiếp thị Toàn cầu (Global MarTech) ước tính đạt 344.8 tỷ USD vào năm 2026.")
doc.add_paragraph("Thị trường có thể phục vụ (SAM): Nhóm doanh nghiệp SME và Marketing Agency quy mô nhỏ tại khu vực Châu Á - Thái Bình Dương (APAC) với quy mô ước tính 25 tỷ USD.")
doc.add_paragraph("Thị trường có thể tiếp cận (SOM): Tập trung vào nhóm 10,000 Agency và SME tiên phong ứng dụng công nghệ trong vòng 2 năm đầu tiên, kỳ vọng mức doanh thu thường niên đạt 5 triệu USD (ARR).")

add_heading('3. KIẾN TRÚC SẢN PHẨM VÀ CÔNG NGHỆ CỐT LÕI (PRODUCT & TECHNOLOGY)', 1)
add_heading('3.1. Hệ Sinh Thái 5 Đặc Vụ Chuyên Trách (The 5-Agent Network)', 2)
doc.add_paragraph("BrandFlow không sử dụng một mô hình duy nhất. Thay vào đó, chúng tôi kiến tạo một hệ thống đa tác vụ, nơi mỗi AI được 'Fine-tune' (tinh chỉnh) cực kỳ chuyên sâu:")
doc.add_paragraph("1. Intake Agent (Đặc vụ Dữ Liệu): Tiếp nhận tài liệu, hồ sơ thô của khách hàng. Phân tách và cấu trúc hóa thành 'DNA Thương Hiệu'.")
doc.add_paragraph("2. Strategy Agent (Đặc vụ Chiến Lược): Trí tuệ cốt lõi đóng vai trò Giám đốc Marketing. Lập phễu người dùng (AIDA), phân bổ kênh truyền thông (Omnichannel) và xây dựng Concept định vị.")
doc.add_paragraph("3. Design Agent (Đặc vụ Nghệ Thuật): Quản trị tính nhất quán của Visual. Đề xuất bảng màu (Color Palette), Typography và tự động xuất các bản Mockup giao diện, sản phẩm.")
doc.add_paragraph("4. Content Agent (Đặc vụ Nội Dung): Kế thừa định vị từ Strategy Agent để sản xuất các thông điệp truyền thông (Key Messages) và kế hoạch nội dung chi tiết.")
doc.add_paragraph("5. Math Engine (Đặc vụ Quản Trị Rủi Ro): Cỗ máy tính toán tài chính trung tâm. Mọi đề xuất tiêu tiền đều phải đi qua chốt chặn này.")

add_heading('3.2. Đột Phá Công Nghệ: Math Engine', 2)
doc.add_paragraph("Khi Strategy Agent yêu cầu ngân sách 15,000 USD cho nền tảng TikTok, Math Engine sẽ tiến hành phân tích dự phóng. Dựa trên thuật toán tính điểm rủi ro, nếu Math Engine nhận thấy tỷ lệ CPA (Cost Per Acquisition) trong ngành đang quá cao dẫn đến biên lợi nhuận ròng (Net Margin) âm, hệ thống sẽ tự động phủ quyết lệnh của Strategy Agent và yêu cầu thiết lập lại phương án khả thi hơn. Tính năng này được chúng tôi đăng ký độc quyền công nghệ bảo chứng an toàn tài chính.")

add_heading('4. MÔ HÌNH KINH DOANH VÀ DỰ PHÓNG TÀI CHÍNH (BUSINESS MODEL)', 1)
add_heading('4.1. Cơ cấu sản phẩm (Tiered SaaS Model)', 2)
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Tên Gói'
hdr_cells[1].text = 'Giá (USD/Tháng)'
hdr_cells[2].text = 'Phân Khúc Khách Hàng'
hdr_cells[3].text = 'Tính Năng Nổi Bật'

row = table.add_row().cells
row[0].text = 'Khởi Điểm (Free)'
row[1].text = '$0'
row[2].text = 'Freelancer, Cá nhân'
row[3].text = 'Sử dụng 1 Agent, giới hạn 10 prompt/ngày, có watermark'

row = table.add_row().cells
row[0].text = 'Nâng Cao (Pro)'
row[1].text = '$19'
row[2].text = 'SME, Marketer độc lập'
row[3].text = 'Full 5 Agents, xuất file PDF White-label, ưu tiên hàng đợi'

row = table.add_row().cells
row[0].text = 'Doanh Nghiệp (Enterprise)'
row[1].text = 'Tùy chỉnh'
row[2].text = 'Tập đoàn, Global Agency'
row[3].text = 'On-premise, Custom LLM Training, SLA 99.9%, hỗ trợ 24/7'

doc.add_paragraph('\n')
add_heading('4.2. Chỉ số Đơn vị Tài chính (Unit Economics)', 2)
doc.add_paragraph("Dựa trên báo cáo kiểm thử 10,000 người dùng thử nghiệm đầu tiên, hệ thống ghi nhận các chỉ số vượt trội:", style='List Bullet')
doc.add_paragraph("Tỷ lệ chuyển đổi (Conversion Rate): 10.0% khách hàng Free chuyển đổi lên gói Pro.", style='List Bullet 2')
doc.add_paragraph("Chi phí thu thập khách hàng (CAC): 50 USD/Khách hàng trả phí.", style='List Bullet 2')
doc.add_paragraph("Giá trị vòng đời (LTV): 266 USD/Khách hàng (Tỷ lệ Churn Rate ổn định ở mức 5%).", style='List Bullet 2')
doc.add_paragraph("Tỷ lệ LTV:CAC = 5.3 : 1 (Chứng minh năng lực sinh lời bền vững ở quy mô lớn).", style='List Bullet 2')

add_heading('5. LỘ TRÌNH PHÁT TRIỂN & KÊU GỌI ĐẦU TƯ (ROADMAP & THE ASK)', 1)
doc.add_paragraph("Dự án đang tiến vào Giai đoạn Tăng tốc (Scale-up Phase) với lộ trình 24 tháng như sau:")
doc.add_paragraph("Tháng 1-6: Hoàn thiện hạ tầng Cloud Server, gia tăng khả năng chịu tải cho mạng lưới AI.")
doc.add_paragraph("Tháng 7-12: Triển khai chiến dịch thâm nhập thị trường (Penetration Pricing), tập trung vào tập khách hàng Agency.")
doc.add_paragraph("Tháng 13-24: Mở rộng hệ sinh thái ứng dụng, phát hành bộ Open API cho đối tác tích hợp.")

doc.add_paragraph("Lời Kêu Gọi Đầu Tư (The Ask): Chúng tôi chào đón các Đối tác Tài chính và Quỹ Đầu tư Mạo hiểm (VCs) đồng hành cùng BrandFlow trong vòng gọi vốn Hạt giống (Seed Round), nhằm mở rộng năng lực xử lý phần cứng GPU và tăng trưởng thị phần tại khu vực APAC.")

os.makedirs('docs', exist_ok=True)
doc.save('docs/BrandFlow_Business_Plan_Detailed.docx')
print('Docx generated successfully!')
